import docx
from docx.shared import Pt


def definir_margens_titulos(caminho_arquivo, margem_esquerda_pt, margem_direita_pt):
    """
    Define a mesma margem esquerda e direita para os estilos de título
    (Título 1, Título 2 e Título 3) em um documento .docx.

    Args:
        caminho_arquivo (str): O caminho para o arquivo .docx.
        margem_esquerda_pt (int): O valor da margem esquerda em pontos (Pt).
        margem_direita_pt (int): O valor da margem direita em pontos (Pt).
    """
    try:
        documento = docx.Document(caminho_arquivo)
        estilos = documento.styles

        # Nomes dos estilos de título que serão modificados
        nomes_estilos_titulo = ["Heading 1", "Heading 2", "Heading 3"]

        for nome_estilo in nomes_estilos_titulo:
            if nome_estilo in estilos:
                estilo = estilos[nome_estilo]
                formato_paragrafo = estilo.paragraph_format
                formato_paragrafo.left_indent = Pt(margem_esquerda_pt)
                formato_paragrafo.right_indent = Pt(margem_direita_pt)
                print(
                    f"Margens do estilo '{nome_estilo}' atualizadas para {margem_esquerda_pt}pt (esquerda) e {margem_direita_pt}pt (direita)."
                )
            else:
                print(
                    f"Aviso: O estilo '{nome_estilo}' não foi encontrado no documento."
                )

        # Salva o documento com as modificações
        documento.save(caminho_arquivo)
        print("\nDocumento salvo com sucesso!")

    except FileNotFoundError:
        print(f"Erro: O arquivo '{caminho_arquivo}' não foi encontrado.")
    except Exception as e:
        print(f"Ocorreu um erro: {e}")


import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH


def centralizar_imagens_na_pagina(caminho_arquivo):
    """
    Centraliza todos os parágrafos que contêm uma imagem em um documento .docx.
    Para um alinhamento "real" com a página, o método zera os recuos do
    parágrafo antes de centralizá-lo.

    IMPORTANTE: Funciona apenas para imagens inseridas "Em Linha com o Texto".
    Para melhor resultado, cada imagem deve estar em seu próprio parágrafo.

    Args:
        caminho_arquivo (str): O caminho para o arquivo .docx.
    """
    try:
        documento = docx.Document(caminho_arquivo)
        paragrafos_centralizados = 0
        print("Analisando o documento para centralizar imagens...")

        for paragrafo in documento.paragraphs:
            # A forma mais robusta de detectar uma imagem é verificar a tag '<w:drawing>'
            # no XML do parágrafo.
            if "<w:drawing>" in paragrafo._p.xml:
                # 1. Zera os recuos esquerdo e direito do parágrafo
                # Isso faz o parágrafo se expandir por toda a largura entre as margens da página
                formato_paragrafo = paragrafo.paragraph_format
                formato_paragrafo.left_indent = Pt(0)
                formato_paragrafo.right_indent = Pt(0)

                # 2. Define o alinhamento do parágrafo como centralizado
                paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

                paragrafos_centralizados += 1
                print(
                    "Parágrafo com imagem encontrado e centralizado em relação às margens."
                )

        if paragrafos_centralizados > 0:
            documento.save(caminho_arquivo)
            print(
                f"\nOperação concluída. {paragrafos_centralizados} parágrafo(s) com imagem foram centralizados."
            )
        else:
            print(
                "\nOperação concluída. Nenhuma imagem 'Em Linha com o Texto' foi encontrada."
            )

    except FileNotFoundError:
        print(f"ERRO: O arquivo '{caminho_arquivo}' não foi encontrado.")
    except Exception as e:
        print(f"Ocorreu um erro inesperado: {e}")


# --- Exemplo de Uso ---
if __name__ == "__main__":
    # Crie um documento de exemplo para o teste
    # doc = docx.Document()
    # doc.add_heading("Este é o Título 1", level=1)
    # doc.add_paragraph("Este é um parágrafo de texto normal.")
    # doc.add_heading("Este é o Título 2", level=2)
    # doc.add_paragraph("Este é outro parágrafo de texto.")
    # doc.add_heading("Este é o Título 3", level=3)
    # doc.add_paragraph("Mais um parágrafo.")
    # doc.save("documento_exemplo.docx")

    caminho_do_documento = r"D:\AT-Victor\Arquivos\Novo-Teste\teste.docx"
    margem_esquerda = 36  # 1 polegada (72 pontos)
    margem_direita = 36  # 1 polegada (72 pontos)

    definir_margens_titulos(caminho_do_documento, margem_esquerda, margem_direita)
    centralizar_imagens_na_pagina(caminho_do_documento)
