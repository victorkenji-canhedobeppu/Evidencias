# src/core/doc_appender.py
# VERSÃO FINAL E DEFINITIVA: Usa win32com para programaticamente criar e aplicar
# o esquema de numeração hierárquico, garantindo que funcione em qualquer documento.

from math import ceil
import os
import fitz
import tempfile
import docx
import pandas as pd
import numpy as np
from docx.shared import Pt
from datetime import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH
import win32com.client as win32
from win32com.client import constants as c
from tkinter import messagebox

from config.settings import COLUMN_RENAME_MAP, DISCIPLINE_SHEET_TYPES, SHEET_COLUMN_MAP


class DocxAppender:
    def __init__(
        self,
        docx_path: str,
        measurement_month: int,
        measurement_year: int,
    ):
        self.docx_path = os.path.abspath(docx_path)
        self.measurement_month = measurement_month
        self.measurement_year = measurement_year
        self.all_excel_data = pd.DataFrame()

    def _find_files(self, folder_path: str):
        if not os.path.isdir(folder_path):
            return []
        supported_files = []
        for f in os.listdir(folder_path):
            if f.lower().endswith((".jpeg", ".jpg", ".png", ".pdf", ".xlsx")):
                supported_files.append(os.path.join(folder_path, f))
        return sorted(supported_files)

    def _apply_heading_numbering(self, word_app, doc):
        """
        Constrói e aplica um esquema de numeração hierárquico aos estilos de Título.
        Esta função é a chave para garantir que a numeração funcione sempre.
        """
        try:
            # Pega na galeria de listas de múltiplos níveis
            gallery = word_app.ListGalleries(c.wdOutlineNumberGallery)

            # Adiciona um novo template de lista ao documento. Isto "reseta" a formatação.
            list_template = doc.ListTemplates.Add(True)

            text_position_pt = 1.5 * 28.3464567

            lvl1 = list_template.ListLevels(1)
            lvl1.NumberFormat = "%1"  # Formato "NúmeroDoNível1.NúmeroDoNível2"
            lvl1.TrailingCharacter = c.wdTrailingTab
            lvl1.NumberStyle = c.wdListNumberStyleArabic
            lvl1.LinkedStyle = "Título 1"
            lvl1.NumberPosition = 0
            lvl1.TextPosition = text_position_pt
            lvl1.TabPosition = text_position_pt

            # Define o Nível 2 para estar ligado ao estilo "Título 2"
            lvl2 = list_template.ListLevels(2)
            lvl2.NumberFormat = "%1.%2"  # Formato "NúmeroDoNível1.NúmeroDoNível2"
            lvl2.TrailingCharacter = c.wdTrailingTab
            lvl2.NumberStyle = c.wdListNumberStyleArabic
            lvl2.LinkedStyle = "Título 2"
            lvl2.NumberPosition = 0
            lvl2.TextPosition = text_position_pt
            lvl2.TabPosition = text_position_pt

            # Define o Nível 3 para estar ligado ao estilo "Título 3"
            lvl3 = list_template.ListLevels(3)
            lvl3.NumberFormat = "%1.%2.%3"  # Formato "N1.N2.N3"
            lvl3.TrailingCharacter = c.wdTrailingTab
            lvl3.NumberStyle = c.wdListNumberStyleArabic
            lvl3.LinkedStyle = "Título 3"
            lvl3.NumberPosition = 0
            lvl3.TextPosition = text_position_pt
            lvl3.TabPosition = text_position_pt

            print("Numeração dos estilos de título foi configurada com sucesso.")
            return True
        except Exception as e:
            print(f"Não foi possível configurar a numeração dos títulos. Erro: {e}")
            messagebox.showerror(
                "Erro de Configuração",
                f"Não foi possível configurar a numeração automática no Word.\n\nDetalhe: {e}",
            )
            return False

    def _convert_pdf_to_png(self, pdf_path):
        """
        Converte a primeira página de um PDF para um ficheiro PNG,
        salvando-o na mesma pasta com o mesmo nome.
        """
        try:
            # Constrói o caminho para o ficheiro de saída
            output_png_path = os.path.splitext(pdf_path)[0] + ".png"

            # Se a imagem já existir, não a converte novamente
            if os.path.exists(output_png_path):
                print(
                    f"Imagem PNG já existe para {os.path.basename(pdf_path)}. A usar a existente."
                )
                return output_png_path

            doc = fitz.open(pdf_path)
            page = doc.load_page(0)

            # Aumenta o zoom para melhor qualidade da imagem (dpi)
            zoom = 2
            matrix = fitz.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=matrix)
            doc.close()

            # Salva a imagem no caminho de saída
            pix.save(output_png_path)
            print(f"PDF convertido para: {output_png_path}")
            return output_png_path

        except Exception as e:
            print(f"Erro ao converter PDF {os.path.basename(pdf_path)}: {e}")
            return None

    def _define_magin_title(self, path, margin_left, margin_right):

        try:
            doc = docx.Document(path)
            styles = doc.styles

            # Nomes dos style de título que serão modificados
            styles_name_title = ["Heading 1", "Heading 2", "Heading 3"]

            for style_name in styles_name_title:
                if style_name in styles:
                    estilo = styles[style_name]
                    para_format = estilo.paragraph_format
                    para_format.left_indent = Pt(margin_left)
                    para_format.right_indent = Pt(margin_right)
                    print(
                        f"Margens do estilo '{style_name}' atualizadas para {margin_left}pt (esquerda) e {margin_right}pt (direita)."
                    )
                else:
                    print(
                        f"Aviso: O estilo '{style_name}' não foi encontrado no documento."
                    )

            # Salva o documento com as modificações
            doc.save(path)
            print("\nDocumento salvo com sucesso!")

        except FileNotFoundError:
            print(f"Erro: O arquivo '{path}' não foi encontrado.")
        except Exception as e:
            print(f"Ocorreu um erro: {e}")

    def _centralize_image(self, caminho_arquivo):
        try:
            doc = docx.Document(caminho_arquivo)
            centralized_paragraphs = 0

            for paragraph in doc.paragraphs:
                # A forma mais robusta de detectar uma imagem é verificar a tag '<w:drawing>'
                # no XML do parágrafo.
                if "<w:drawing>" in paragraph._p.xml:
                    # 1. Zera os recuos esquerdo e direito do parágrafo
                    # Isso faz o parágrafo se expandir por toda a largura entre as margens da página
                    paragraph_format = paragraph.paragraph_format
                    paragraph_format.left_indent = Pt(0)
                    paragraph_format.right_indent = Pt(0)

                    # 2. Define o alinhamento do parágrafo como centralizado
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    centralized_paragraphs += 1

            if centralized_paragraphs > 0:
                doc.save(caminho_arquivo)
                print(
                    f"\nOperação concluída. {centralized_paragraphs} parágrafo(s) com imagem foram centralizados."
                )
            else:
                print(
                    "\nOperação concluída. Nenhuma imagem 'Em Linha com o Texto' foi encontrada."
                )

        except FileNotFoundError:
            print(f"ERRO: O arquivo '{caminho_arquivo}' não foi encontrado.")
        except Exception as e:
            print(f"Ocorreu um erro inesperado: {e}")

    def _process_excel_file(
        self,
        excel_path: str,
        measurement_month: int,
        measurement_year: int,
        discipline: str,
    ) -> dict:

        processed_data = {}
        valid_discipline = DISCIPLINE_SHEET_TYPES.get(discipline, [])
        print(f"Processando abas do tipo: '{valid_discipline}'")

        try:
            excel_file = pd.ExcelFile(excel_path)
            for sheet_name in excel_file.sheet_names:
                if sheet_name in valid_discipline:
                    if sheet_name.upper() in SHEET_COLUMN_MAP:
                        header_df = pd.read_excel(
                            excel_file, sheet_name=sheet_name, header=None, nrows=7
                        )
                        row6, row7 = header_df.iloc[5], header_df.iloc[6]
                        final_headers_raw = [
                            h7 if pd.notna(h7) else h6 for h6, h7 in zip(row6, row7)
                        ]
                        counts = {}
                        final_headers = []
                        for item in final_headers_raw:
                            item_str = (
                                str(item)
                                if pd.notna(item)
                                else f"Unnamed_{len(final_headers)}"
                            )
                            counts[item_str] = counts.get(item_str, 0) + 1
                            if counts[item_str] > 1:
                                final_headers.append(f"{item_str}_{counts[item_str]}")
                            else:
                                final_headers.append(item_str)

                        data_df = pd.read_excel(
                            excel_file, sheet_name=sheet_name, header=None, skiprows=7
                        )
                        data_df.columns = final_headers[: len(data_df.columns)]

                        total_row_indices = np.where(
                            data_df.apply(
                                lambda row: row.astype(str)
                                .str.contains("TOTAL", case=False, na=False)
                                .any(),
                                axis=1,
                            )
                        )
                        if total_row_indices[0].size > 0:
                            data_df = data_df.iloc[: total_row_indices[0][0]]

                        if sheet_name.upper() in (
                            "TRADO",
                            "SHELBY",
                            "DENISON",
                            "BLOCO",
                            "POÇO",
                        ):
                            # --- ALTERAÇÃO PRINCIPAL AQUI ---
                            for col_name in data_df.columns:
                                if str(col_name).upper().startswith("MEDIÇÃO"):
                                    col_idx = data_df.columns.get_loc(col_name)
                                    if col_idx > 0:
                                        # Pega o nome da coluna à esquerda da data
                                        left_col_name = data_df.columns[col_idx - 1]

                                        # Define as colunas de identificação que não devem ser apagadas
                                        protected_cols = [
                                            "AMOSTRAGEM SHELBY",
                                            "AMOSTRAGEM DENISON",
                                            "BLOCO INDEFORMADO",
                                        ]

                                        # Aplica a limpeza APENAS SE a coluna à esquerda não for uma das protegidas
                                        if (
                                            str(left_col_name).upper()
                                            not in protected_cols
                                        ):
                                            dt_col = pd.to_datetime(
                                                data_df[col_name], errors="coerce"
                                            )
                                            null_rows = (
                                                dt_col.dt.year != measurement_year
                                            ) | (dt_col.dt.month != measurement_month)
                                            data_df.iloc[null_rows, col_idx - 1] = (
                                                np.nan
                                            )
                            # --- FIM DA ALTERAÇÃO ---

                            any_medicao_valid = pd.Series(False, index=data_df.index)
                            for col_name in data_df.columns:
                                if str(col_name).upper().startswith("MEDIÇÃO"):
                                    dt_col = pd.to_datetime(
                                        data_df[col_name], errors="coerce"
                                    )
                                    valid_rows_mask = (
                                        dt_col.dt.year == measurement_year
                                    ) & (dt_col.dt.month == measurement_month)
                                    any_medicao_valid = (
                                        any_medicao_valid
                                        | valid_rows_mask.fillna(False)
                                    )
                            data_df = data_df[any_medicao_valid]
                        else:
                            if "MEDIÇÃO" in data_df.columns:
                                medicao_col_as_datetime = pd.to_datetime(
                                    data_df["MEDIÇÃO"], errors="coerce"
                                )
                                mask = (
                                    medicao_col_as_datetime.dt.year == measurement_year
                                ) & (
                                    medicao_col_as_datetime.dt.month
                                    == measurement_month
                                )
                                data_df = data_df[mask]

                        if not data_df.empty:
                            final_cols = [
                                col
                                for col in SHEET_COLUMN_MAP[sheet_name.upper()]
                                if col in data_df.columns
                            ]
                            df_to_add = data_df[final_cols].dropna(how="all")
                            df_to_add.rename(columns=COLUMN_RENAME_MAP, inplace=True)

                            if sheet_name.upper() in ("SHELBY", "BLOCO", "DENISON"):
                                source_column_name = "Umidade"
                                new_column_name = "Número de camadas"

                                if source_column_name in df_to_add.columns:
                                    df_to_add[new_column_name] = df_to_add[
                                        source_column_name
                                    ]
                                    cols = df_to_add.columns.tolist()
                                    umidade_index = cols.index(source_column_name)
                                    cols.pop(cols.index(new_column_name))
                                    cols.insert(umidade_index, new_column_name)
                                    df_to_add = df_to_add[cols]
                                    print(
                                        f"Coluna '{new_column_name}' adicionada e reordenada na aba '{sheet_name}'."
                                    )

                            processed_data[sheet_name] = df_to_add
        except Exception as e:
            print(f"Erro ao processar o arquivo Excel '{excel_path}': {e}")

        return processed_data

    def _add_dataframe_as_table(self, doc, df: pd.DataFrame, table_range, word_app):
        """Adiciona um DataFrame como uma tabela em um range específico do documento."""
        if df.empty:
            return

        # Importa as constantes do Word para uso local
        from win32com.client import constants as c

        table = doc.Tables.Add(
            Range=table_range, NumRows=df.shape[0] + 1, NumColumns=df.shape[1]
        )
        table.Select()

        # 2. Aplica a formatação de parágrafo à SELEÇÃO atual (a tabela in
        word_app.Selection.ParagraphFormat.Alignment = c.wdAlignParagraphCenter
        table.Style = "Tabela com grade"
        table.Range.ParagraphFormat.Alignment = c.wdAlignParagraphCenter
        table.Borders.Enable = True

        # Formatação do Cabeçalho
        header_row = table.Rows(1)
        header_row.HeadingFormat = True
        header_row.Shading.BackgroundPatternColor = c.wdColorDarkBlue
        header_font = header_row.Range.Font
        header_font.Size = 7
        header_font.Name = "Arial"
        header_font.ColorIndex = c.wdWhite
        header_font.Bold = True
        header_row.Range.ParagraphFormat.Alignment = c.wdAlignParagraphCenter

        for j, col_name in enumerate(df.columns):
            cell = table.Cell(Row=1, Column=j + 1)
            cell.Range.Text = str(col_name)
            cell.WordWrap = False
            cell.VerticalAlignment = c.wdCellAlignVerticalCenter

        # Preenchimento dos dados
        for i, row_data in enumerate(df.itertuples(index=False)):
            for j, val in enumerate(row_data):
                cell = table.Cell(i + 2, j + 1)
                cell_text = ""
                if pd.notna(val):
                    if isinstance(val, (datetime, pd.Timestamp)):
                        cell_text = val.strftime("%d/%m/%Y")
                    else:
                        cell_text = str(val)

                cell.Range.Text = cell_text
                cell.Range.Font.Size = 7
                cell.Range.Font.Name = "Arial"
                cell.WordWrap = False
                cell.VerticalAlignment = c.wdCellAlignVerticalCenter

        try:
            table.AutoFitBehavior(c.wdAutoFitContent)
        except Exception as e:
            print(f"Não foi possível aplicar o AutoFit: {e}")

        print(f"Tabela com {df.shape[0]} linhas inserida.")

    def append_measurement(
        self,
        disciplines: list,
        user_texts: dict,
        base_evidence_path: str,
        measurement_month: int,
        measurement_year: int,
    ):
        sondagens_str = "Sondagens"
        ensaios_str = "Ensaios Especiais"

        if sondagens_str in disciplines and ensaios_str in disciplines:
            idx_sondagens = disciplines.index(sondagens_str)
            idx_ensaios = disciplines.index(ensaios_str)

            # Se "Sondagens" estiver depois de "Ensaios Especiais", corrige a ordem
            if idx_sondagens > idx_ensaios:
                # Remove 'Sondagens' da sua posição atual
                disciplines.remove(sondagens_str)
                # Recalcula o índice de 'Ensaios' (pode ter mudado) e insere 'Sondagens' antes
                new_idx_ensaios = disciplines.index(ensaios_str)
                disciplines.insert(new_idx_ensaios, sondagens_str)
                print(f"Ordem das disciplinas ajustada: {disciplines}")

        word_app = None
        doc = None
        try:
            word_app = win32.Dispatch("Word.Application")
            word_app.Visible = False
            doc = word_app.Documents.Open(self.docx_path)

            # Importa as constantes do Word para uso local
            from win32com.client import constants as c

            if not self._apply_heading_numbering(word_app, doc):
                return

            # Move para o final do documento para começar a adicionar conteúdo
            selection = word_app.Selection
            selection.EndKey(Unit=c.wdStory)
            selection.TypeParagraph()  # Garante que estamos em um novo parágrafo

            # --- Título Principal da Medição ---
            mes_formatado = str(self.measurement_month).zfill(2)
            selection.Style = "Título 2"
            selection.Font.Name = "Arial"
            selection.TypeText(
                Text=f"MEDIÇÃO ({mes_formatado}/{self.measurement_year})"
            )
            selection.TypeParagraph()

            selection.Style = "Título 3"
            selection.Font.Name = "Arial"
            selection.TypeText(Text="Projeto")
            selection.TypeParagraph()

            user_text = user_texts.get("Projeto", "")
            if user_text:
                selection.Style = "Normal"
                selection.Font.Name = "Arial"
                selection.TypeText(Text=user_text)
                selection.TypeParagraph()

            # --- Loop Principal por Disciplina ---
            for discipline in disciplines:
                selection.Style = "Título 3"
                selection.Font.Name = "Arial"
                selection.TypeText(Text=discipline)
                selection.TypeParagraph()

                user_text = user_texts.get(discipline, "")
                if user_text:
                    selection.Style = "Normal"
                    selection.Font.Name = "Arial"
                    selection.TypeText(Text=user_text)
                    selection.TypeParagraph()

                # --- Processamento de Ficheiros (Imagens, PDFs) ---
                discipline_folder_path = os.path.join(base_evidence_path, discipline)
                files_to_add = self._find_files(discipline_folder_path)
                excel_files = [
                    f for f in files_to_add if f.lower().endswith((".xlsx", ".xls"))
                ]

                image_files = [
                    f
                    for f in files_to_add
                    if f.lower().endswith((".jpeg", ".jpg", ".png", ".pdf"))
                ]
                if not image_files and not excel_files:
                    selection.TypeText(
                        Text="Nenhuma evidência encontrada para esta disciplina."
                    )
                    selection.Font.Name = "Arial"
                    selection.TypeParagraph()

                for file_path in image_files:
                    abs_path = os.path.abspath(file_path)
                    img_path_to_add = abs_path
                    if file_path.lower().endswith(".pdf"):
                        img_path_to_add = self._convert_pdf_to_png(abs_path)

                    if img_path_to_add:
                        selection.InlineShapes.AddPicture(
                            FileName=img_path_to_add,
                            LinkToFile=False,
                            SaveWithDocument=True,
                        )
                        selection.TypeParagraph()

                # --- Processamento de Tabelas Excel (Lógica Central) ---
                if excel_files:
                    for excel_path in excel_files:
                        try:
                            all_excel_data = self._process_excel_file(
                                excel_path,
                                measurement_month,
                                measurement_year,
                                discipline,
                            )
                            WIDE_TABLE_SHEETS = [
                                "TRADO",
                                "SHELBY",
                                "DENISON",
                                "BLOCO",
                                "POÇO",
                            ]

                            for sheet_name, final_df in all_excel_data.items():

                                final_df = final_df.dropna(axis=1, how="all").dropna(
                                    axis=0, how="all"
                                )

                                # 2. Pula para a próxima aba se o DataFrame ficar vazio após a limpeza.
                                if final_df.empty:
                                    continue

                                selection.ParagraphFormat.Alignment = (
                                    c.wdAlignParagraphCenter
                                )
                                selection.Font.Bold = True
                                selection.TypeText(f"Tabela: {sheet_name}")
                                selection.Font.Bold = False

                                # 2. Pula para o próximo parágrafo (que também será centralizado)
                                selection.TypeParagraph()

                                # Define o range onde a tabela será inserida
                                current_range = selection.Range

                                if sheet_name.upper() in WIDE_TABLE_SHEETS:
                                    key_cols = ["Sondagem"]
                                    actual_keys = [
                                        c for c in key_cols if c in final_df.columns
                                    ]
                                    data_cols = [
                                        c
                                        for c in final_df.columns
                                        if c not in actual_keys
                                    ]
                                    MAX_DATA_COLS = 9
                                    num_chunks = ceil(len(data_cols) / MAX_DATA_COLS)

                                    for i in range(num_chunks):
                                        chunk_cols = data_cols[
                                            i * MAX_DATA_COLS : (i + 1) * MAX_DATA_COLS
                                        ]
                                        table_df = final_df[actual_keys + chunk_cols]

                                        # Adiciona a tabela no range atual e move o cursor para depois dela
                                        self._add_dataframe_as_table(
                                            doc, table_df, current_range, word_app
                                        )
                                        selection.EndKey(
                                            Unit=c.wdStory
                                        )  # Move para o fim para continuar
                                        selection.TypeParagraph()
                                        current_range = selection.Range

                                else:
                                    # Adiciona a tabela no range atual e move o cursor para depois dela
                                    self._add_dataframe_as_table(
                                        doc, final_df, current_range, word_app
                                    )
                                    selection.EndKey(
                                        Unit=c.wdStory
                                    )  # Move para o fim para continuar
                                    selection.TypeParagraph()

                        except Exception as e:
                            print(
                                f"Erro CRÍTICO ao processar tabelas do Excel '{excel_path}': {e}"
                            )

            # --- Finalização ---
            doc.Fields.Update()
            doc.Save()
            messagebox.showinfo(
                "Sucesso", f"Conteúdo adicionado com sucesso em:\n{self.docx_path}"
            )

        except Exception as e:
            messagebox.showerror(
                "Erro Fatal", f"Ocorreu um erro fatal ao interagir com o Word:\n{e}"
            )
        finally:
            if doc:
                doc.Close(SaveChanges=0)
            if word_app:
                word_app.Quit()

            # Funções de formatação final
            self._define_magin_title(self.docx_path, 36, 36)
            self._centralize_image(self.docx_path)
